# src/reporting.py
from __future__ import annotations
from typing import Dict, Iterable, Mapping, Tuple
from copy import deepcopy
from pathlib import Path

import numpy as np
import pandas as pd

# Optional dependency (only used by export_table_to_docx)
try:
    from docx import Document
    from docx.shared import Pt
    _HAS_DOCX = True
except Exception:
    _HAS_DOCX = False


"""
reporting.py

Tabular summaries for cohort coverage, window counts, split summaries, and
model hyperparameter tables, plus (optional) Word export.

Public API:
- build_union_span(...) / union_span_from_daily(...)
- presence_days(...)
- summarise_stream(...)
- coverage_for_columns(...)
- summary_from_per_user(...)
- summarise_windows_dict(...)
- build_split_summary(...)
- ids_from_Xy(...)
- xgb_params_table(...), lr_params_table(...)
- export_table_to_docx(...)

Notes:
- This module contains *no plotting*. Visuals live in src.viz.*
- Word export is optional; will raise a clear error if python-docx is missing.
"""


# ---------------------------- Cohort span utilities ----------------------------

def _span(df: pd.DataFrame, user_col: str = "user_key", date_col: str = "date") -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame(columns=[user_col, "min_date", "max_date"]).set_index(user_col)
    g = df.groupby(user_col)[date_col]
    return g.min().to_frame("min_date").join(g.max().to_frame("max_date"))


def build_union_span(
    hr_df: pd.DataFrame | None,
    inh_df: pd.DataFrame | None,
    dq_df: pd.DataFrame | None,
    user_col: str = "user_key",
    date_col: str = "date",
) -> pd.DataFrame:
    """
    Union over min/max dates from multiple daily-level dataframes (HR, inhaler, questionnaire).
    Returns a per-user table with span_days.
    """
    span = (
        pd.concat([_span(hr_df, user_col, date_col),
                   _span(inh_df, user_col, date_col),
                   _span(dq_df, user_col, date_col)])
          .groupby(level=0)
          .agg({"min_date": "min", "max_date": "max"})
          .dropna()
    )
    span["span_days"] = (span["max_date"] - span["min_date"] + 1).astype("Int64")
    return span


def union_span_from_daily(
    dfs: Iterable[pd.DataFrame],
    user_col: str = "user_key",
    date_col: str = "date",
) -> pd.DataFrame:
    """Generalised version of build_union_span for any iterable of daily dataframes."""
    spans = []
    for df in dfs:
        if df is None or df.empty:
            continue
        g = df.groupby(user_col)[date_col]
        spans.append(g.min().to_frame("min_date").join(g.max().to_frame("max_date")))
    if not spans:
        return pd.DataFrame(columns=[user_col, "min_date", "max_date", "span_days"]).set_index(user_col)
    span = (
        pd.concat(spans)
          .groupby(level=0)
          .agg({"min_date": "min", "max_date": "max"})
          .dropna()
    )
    span["span_days"] = (span["max_date"] - span["min_date"] + 1).astype("Int64")
    return span


def presence_days(df: pd.DataFrame | None, user_col: str = "user_key", date_col: str = "date") -> pd.Series:
    """Number of distinct observed days per user for a daily dataframe."""
    if df is None or df.empty:
        return pd.Series(dtype="Int64")
    d = df[[user_col, date_col]].dropna().drop_duplicates()
    return d.groupby(user_col)[date_col].nunique().astype("Int64")


def summarise_stream(name: str, present: pd.Series, span_tbl: pd.DataFrame) -> Dict:
    """
    Cohort-level summary for a single stream given per-user 'present days' and span table.
    Returns a dict suitable for DataFrame([...]).
    """
    if present is None or present.empty or span_tbl is None or span_tbl.empty:
        return {"stream": name, "participants": 0, "total_days": 0, "median_days": "NA",
                "IQR_low": "NA", "IQR_high": "NA", "min_days": "NA", "max_days": "NA",
                "mean_%_completeness": 0.0}
    tmp = span_tbl[["span_days"]].join(present.rename("present_days"), how="left").fillna(0)
    tmp["present_days"] = tmp["present_days"].astype("Int64")
    with np.errstate(divide="ignore", invalid="ignore"):
        comp = (tmp["present_days"] / tmp["span_days"]).replace([np.inf, -np.inf], np.nan)
    return {
        "stream": name,
        "participants": int((tmp["present_days"] > 0).sum()),
        "total_days": int(tmp["present_days"].sum()),
        "median_days": int(tmp["present_days"].median()),
        "IQR_low": int(tmp["present_days"].quantile(0.25)),
        "IQR_high": int(tmp["present_days"].quantile(0.75)),
        "min_days": int(tmp["present_days"].min()),
        "max_days": int(tmp["present_days"].max()),
        "mean_%_completeness": round(float(np.nanmean(comp) * 100), 1),
    }


# ------------------------------- Coverage & stats ------------------------------

def coverage_for_columns(
    df: pd.DataFrame | None,
    cols: Iterable[str],
    span_tbl: pd.DataFrame,
    user_col: str = "user_key",
    date_col: str = "date",
) -> float:
    """Mean % coverage across users for any subset of columns (days with any non-null)."""
    if df is None or df.empty:
        return np.nan
    cols = [c for c in (cols or []) if c in df.columns]
    if not cols:
        return np.nan
    d = df[[user_col, date_col] + cols].copy()
    d["present"] = d[cols].notna().any(axis=1)
    present = (d.loc[d["present"], [user_col, date_col]]
                 .drop_duplicates()
                 .groupby(user_col)[date_col].nunique())
    tmp = span_tbl[["span_days"]].join(present.rename("present_days"), how="left").fillna(0)
    with np.errstate(divide="ignore", invalid="ignore"):
        comp = (tmp["present_days"] / tmp["span_days"]).replace([np.inf, -np.inf], np.nan)
    return float(np.nanmean(comp) * 100)


def summary_from_per_user(values: pd.Series) -> dict:
    """Median/IQR/min/max from per-user values; returns NaNs if empty."""
    v = pd.to_numeric(values, errors="coerce").dropna()
    if v.empty:
        return {"median": np.nan, "IQR_low": np.nan, "IQR_high": np.nan, "min": np.nan, "max": np.nan}
    return {
        "median": float(np.median(v)),
        "IQR_low": float(np.percentile(v, 25)),
        "IQR_high": float(np.percentile(v, 75)),
        "min": float(np.min(v)),
        "max": float(np.max(v)),
    }


# --------------------------- Window / split summaries --------------------------

def summarise_windows_dict(
    obj_dict: dict,
    *,
    user_col: str = "user_key",
    date_col: str = "date",
) -> Tuple[pd.DataFrame, pd.DataFrame]:
    """
    Build (i) per-user wide table of counts by split (+ date range) and
        (ii) compact across-user summary table (median/IQR/min/max) for each split & total.
    """
    frames = []
    for split in ("train", "val", "test"):
        df = obj_dict.get(split, pd.DataFrame())
        if df is None or df.empty:
            continue
        tmp = df[[user_col, date_col]].copy()
        tmp["split"] = split
        frames.append(tmp)
    if not frames:
        raise ValueError("Nothing to summarise.")
    full = pd.concat(frames, ignore_index=True)

    per_user = (
        full.groupby([user_col, "split"]).size().unstack(fill_value=0)
            .reindex(columns=["train", "val", "test"], fill_value=0)
    )
    per_user["total_windows"] = per_user.sum(axis=1)
    drange = full.groupby(user_col)[date_col].agg(["min", "max"])
    per_user = per_user.join(drange).reset_index()

    def _stats(s: pd.Series) -> pd.Series:
        s = pd.to_numeric(s, errors="coerce").dropna()
        s = s[s > 0]
        if s.empty:
            return pd.Series({"median": 0, "IQR low": 0, "IQR high": 0, "min": 0, "max": 0})
        return pd.Series({
            "median":   int(np.median(s)),
            "IQR low":  int(np.percentile(s, 25)),
            "IQR high": int(np.percentile(s, 75)),
            "min":      int(np.min(s)),
            "max":      int(np.max(s)),
        })

    summary = pd.concat({
        "train_windows": _stats(per_user["train"]),
        "val_windows":   _stats(per_user["val"]),
        "test_windows":  _stats(per_user["test"]),
        "total_windows": _stats(per_user["total_windows"]),
    }, axis=1).T.reset_index().rename(columns={"index": "metric"})

    return per_user, summary


def build_split_summary(obj_dict: dict, *, user_col: str = "user_key") -> pd.DataFrame:
    """Return a small table with users (n) and windows (n) per split."""
    rows = []
    for split in ("train", "val", "test"):
        df = obj_dict.get(split, pd.DataFrame())
        n_users   = int(df[user_col].nunique()) if (isinstance(df, pd.DataFrame) and not df.empty and user_col in df) else 0
        n_windows = int(len(df)) if isinstance(df, pd.DataFrame) else 0
        rows.append({"split": split.capitalize(), "users (n)": n_users, "windows (n)": n_windows})
    return pd.DataFrame(rows, columns=["split", "users (n)", "windows (n)"])


def ids_from_Xy(Xy: dict) -> dict:
    """Convert Xy[split]['ids'] into a {split: DataFrame[user_key, date]} dictionary."""
    out = {}
    for split in ("train", "val", "test"):
        ids = Xy.get(split, {}).get("ids", pd.DataFrame())
        if ids is None or ids.empty:
            out[split] = pd.DataFrame(columns=["user_key", "date"])
        else:
            df = ids.rename(columns={"t_date": "date"})
            out[split] = df[["user_key", "date"]].copy()
    return out


# -------------------------- Hyperparameter summary tables ----------------------

def xgb_params_table(best_dict: dict, _unused: dict | None = None) -> pd.DataFrame:
    """VAL-only hyperparameter table for XGBoost (no TEST metrics)."""
    p = deepcopy(best_dict.get("params", {}))
    rows = [
        ("max_depth",            p.get("max_depth")),
        ("learning_rate (eta)",  p.get("eta")),
        ("subsample",            p.get("subsample")),
        ("colsample_bytree",     p.get("colsample_bytree")),
        ("lambda (L2)",          p.get("lambda")),
        ("alpha (L1)",           p.get("alpha")),
        ("n_rounds",             int(best_dict.get("best_iteration", -1)) + 1),
        ("VAL AUPRC (tuning metric)", round(float(best_dict.get("val_pr_auc", float("nan"))), 3)),
    ]
    return pd.DataFrame(rows, columns=["XGBoost", "value"])


def lr_params_table(best_dict: dict, _unused: dict | None = None) -> pd.DataFrame:
    """VAL-only hyperparameter table for Logistic Regression (no TEST metrics)."""
    p = deepcopy(best_dict.get("params", {}))
    rows = [
        ("penalty",      p.get("penalty")),
        ("C",            float(p.get("C")) if p.get("C") is not None else None),
        ("class_weight", p.get("class_weight")),
        ("solver",       p.get("solver")),
        ("imputer",      "median"),
    ]
    return pd.DataFrame(rows, columns=["Logistic Regression", "value"])


# -------------------------------- Word export ---------------------------------

def export_table_to_docx(
    df: pd.DataFrame,
    filename: str | Path,
    title: str,
    *,
    font: str = "Arial",
    size_pt: int = 10,
):
    """Export a DataFrame to a .docx table (requires python-docx)."""
    if not _HAS_DOCX:
        raise ImportError("python-docx is not installed; install it to use export_table_to_docx().")

    filename = Path(filename)
    doc = Document()
    doc.add_heading(title, level=1)

    tbl = doc.add_table(rows=1, cols=len(df.columns))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, col in enumerate(df.columns):
        run = hdr[i].paragraphs[0].add_run(str(col)); run.font.name = font; run.font.size = Pt(size_pt)

    for row in df.itertuples(index=False):
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            run = cells[i].paragraphs[0].add_run("" if pd.isna(val) else str(val))
            run.font.name = font; run.font.size = Pt(size_pt)

    doc.save(str(filename))
    return str(filename)


__all__ = [
    "build_union_span", "union_span_from_daily", "presence_days", "summarise_stream",
    "coverage_for_columns", "summary_from_per_user",
    "summarise_windows_dict", "build_split_summary", "ids_from_Xy",
    "xgb_params_table", "lr_params_table",
    "export_table_to_docx",
]
